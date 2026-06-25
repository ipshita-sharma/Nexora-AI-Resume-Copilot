from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def export_resume_docx(
    content: str,
    file_path: str
):
    doc = Document()
    content = content.replace(
    "| SUMMARY",
    "\n\nSUMMARY"
)

    content = content.replace(
    "| EDUCATION",
    "\n\nEDUCATION"
)

    content = content.replace(
    "| EXPERIENCE",
    "\n\nEXPERIENCE"
)

    content = content.replace(
    "| PROJECTS",
    "\n\nPROJECTS"
)

    content = content.replace(
    "| TECHNICAL SKILLS",
    "\n\nTECHNICAL SKILLS"
)

    content = content.replace(
    "| CERTIFICATIONS",
    "\n\nCERTIFICATIONS"
)

    section = doc.sections[0]

    section.top_margin = Pt(40)
    section.bottom_margin = Pt(40)
    section.left_margin = Pt(45)
    section.right_margin = Pt(45)

    lines = [
        line.strip()
        for line in content.split("\n")
        if line.strip()
    ]

    print("==== EXPORTED LINES ====")

    for idx, line in enumerate(lines[:15]):
     print(idx, repr(line))

    main_headings = [
    "SUMMARY",
    "EDUCATION",
    "EXPERIENCE",
    "PROJECTS",
    "TECHNICAL SKILLS",
    "CERTIFICATIONS",
    "ACHIEVEMENTS",
    "LEADERSHIP",
    "EXTRACURRICULAR",
    "PUBLICATIONS"
]
    current_section = None
    
    print("==== EXPORTED LINES ====")

    for i, line in enumerate(lines[:20]):
     print(i, repr(line))

    for index, line in enumerate(lines):
        if index == 0:

           name = doc.add_paragraph()

           name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

           run = name.add_run(line)

           run.bold = True
           run.font.size = Pt(16)

           continue
        
        if index == 1:

            contact = doc.add_paragraph()

            contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = contact.add_run(line)

            run.font.size = Pt(10)

            continue

        if index == 2 and "SUMMARY" not in line:

            contact = doc.add_paragraph()

            contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = contact.add_run(line)

            run.font.size = Pt(10)

            continue

        is_heading = line.strip() in main_headings

        is_bullet = (
            line.startswith("-")
            or line.startswith("•")
        )

        is_education_info = (
    "CGPA" in line.upper()
    or "GPA" in line.upper()
) 
        is_skills_content = (
    ":" in line
    and len(line) < 40
    and not is_heading
    and not is_bullet
)
        is_date_line = (
    any(month in line for month in [
        "Jan", "Feb", "Mar", "Apr",
        "May", "Jun", "Jul", "Aug",
        "Sep", "Oct", "Nov", "Dec"
    ])
    and "-" in line
)
# HEADINGS

        if is_heading:
            current_section = line

            heading = doc.add_paragraph()

            run = heading.add_run(line)

            run.bold = True
            run.font.size = Pt(14)

            heading.space_before = Pt(8)
            heading.space_after = Pt(2)

            pPr = heading._element.get_or_add_pPr()

            pBdr = OxmlElement("w:pBdr")

            bottom = OxmlElement("w:bottom")

            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "4F81BD")

            pBdr.append(bottom)

            pPr.append(pBdr)

            continue

        # BULLETS

        if is_bullet:

            cleaned = (
                line.replace("-", "")
                .replace("•", "")
                .strip()
            )

            bullet = doc.add_paragraph(
                style="List Bullet"
            )

            bullet.add_run(cleaned)

            bullet.space_after = Pt(2)

            continue

        if is_education_info:

           edu = doc.add_paragraph()

           run = edu.add_run(line)

           run.font.size = Pt(11)

           edu.space_after = Pt(1)

           continue
        
        if is_date_line:

            date = doc.add_paragraph()

            run = date.add_run(line)

            run.font.size = Pt(10)
            run.italic = True

            date.space_after = Pt(2)

            continue

        is_project_title = (
    (
        "-" in line
        or "–" in line
        or "System" in line
        or "Platform" in line
        or "Application" in line
        or "Intern" in line
        or "Contributor" in line
        or "Developer" in line
    )
    and len(line) < 100
    and not is_bullet
    and not is_heading
    and not is_date_line
    and current_section != "TECHNICAL SKILLS"
)
        
        if is_project_title:

            project = doc.add_paragraph()

            run = project.add_run(line)

            run.bold = True
            run.font.size = Pt(12)

            project.space_before = Pt(4)
            project.space_after = Pt(1)

            continue

        is_tech_stack = (
    "•" in line
    and len(line) < 70
    and not is_heading
    and not is_bullet
)
        if is_tech_stack:

           tech = doc.add_paragraph()

           run = tech.add_run(line)

           run.italic = True
           run.font.size = Pt(10)

           tech.space_after = Pt(2)

           continue

        if is_skills_content:

           skills = doc.add_paragraph()

           run = skills.add_run(line)

           run.font.size = Pt(10.5)

           skills.space_after = Pt(1)

           continue
        # NORMAL TEXT

        paragraph = doc.add_paragraph()

        run = paragraph.add_run(line)

        run.font.size = Pt(11)

        paragraph.space_after = Pt(1)

    doc.save(file_path)